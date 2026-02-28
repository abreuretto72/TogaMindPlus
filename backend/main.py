import os
import sys
import json
from fastapi import FastAPI, Header, Body, Form, HTTPException, UploadFile, File, Request
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from dotenv import load_dotenv
from pydantic import BaseModel, Field
from typing import Optional, List
from colorama import init, Fore
import sqlite3

import time
from toga_logger_service import log_info, log_warning, registrar_erro_com_trace

# Custom Toga modules (Lazy Loading)
_rag_manager = None
_vector_engine = None

def get_rag_manager():
    global _rag_manager
    if _rag_manager is None:
        from toga_rag_manager import TogaRAGManager
        _rag_manager = TogaRAGManager()
    return _rag_manager

def get_vector_engine():
    global _vector_engine
    if _vector_engine is None:
        from toga_vector_engine import TogaVectorEngine
        _vector_engine = TogaVectorEngine()
    return _vector_engine

import requests
import base64
import logging
from logging.handlers import RotatingFileHandler

def get_base_dir_for_log():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.getcwd()

log_file_path = os.path.join(get_base_dir_for_log(), "TogaEngine_audit.log")

logging.basicConfig(
    level=logging.INFO, 
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        RotatingFileHandler(log_file_path, maxBytes=5*1024*1024, backupCount=2, encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# PDF Generator lazy loaded no endpoint correspondenteinit(autoreset=True)

PROMPT_MESTRE_BASE = """
Você é o motor de inteligência artificial TogaMind+, operando como Assessor Jurídico de Gabinete de Alta Performance. Sua função é analisar autos processuais com profundidade acadêmica e precisão técnica absoluta.

DIRETRIZES DE INTEGRIDADE (CRITICAL):
VERDADE ABSOLUTA: Você está proibido de inventar fatos, jurisprudências ou números de processos. Se a resposta não estiver nos dados recuperados (RAG), responda: "Informação não localizada na base de dados consultada".
CITAÇÃO DE FONTES: Cada parágrafo de análise deve conter a citação direta da fonte (ex: "Art. 5º, LVII, CF/88" ou "Página 12 do arquivo Petição_Inicial.pdf").
PROIBIÇÃO DE ALUCINAÇÃO: Jamais complete lacunas de informação com suposições. Se houver ambiguidade, aponte-a como uma dúvida jurídica a ser sanada pelo Magistrado.

CONTROLE DE QUALIDADE FINAL:
Antes de exibir o resultado, realize uma auto-auditoria: "Este dado existe fisicamente nos documentos ou na minha base de treinamento verificável?". Se a resposta for 'não', remova a informação.
"""

# -------------------------------------------------------------------
# VALIDAÇÃO DE AMBIENTE (PROTOCOLO 2026)
# -------------------------------------------------------------------
def validar_ambiente_local():
    """Verifica se os arquivos vitais existem na raiz antes de subir o servidor"""
    erros = []
    raiz = os.getcwd()

    if not os.path.exists(os.path.join(raiz, ".env")):
        erros.append("ERRO: Arquivo '.env' nao encontrado na raiz. Insira sua GEMINI_API_KEY.")

    if not os.path.exists(os.path.join(raiz, "config.json")):
        erros.append("ERRO: Arquivo 'config.json' nao encontrado na raiz.")

# Pasta build/web removida da verificação pois o app é Desktop nativo no Flutter 
    # web_dir = os.path.join(getattr(sys, '_MEIPASS', raiz), "build", "web")
    # if not os.path.exists(web_dir) and not os.path.exists(os.path.join(raiz, "build", "web")):
    #    erros.append("ERRO: Pasta 'build/web' nao encontrada no executavel nem nativamente.")

    if erros:
        print(Fore.RED + "\n" + "="*50)
        print(Fore.RED + "   FALHA NA INICIALIZACAO DO TOGAMIND+")
        print(Fore.RED + "="*50)
        for erro in erros:
            print(Fore.YELLOW + f" > {erro}")
        print(Fore.RED + "="*50 + "\n")
        
        input("Pressione ENTER para sair...")
        sys.exit(1)
    
    print(Fore.GREEN + "[OK] Integridade do sistema verificada. Iniciando Engine...")

# Executa a validação antes de tudo
validar_ambiente_local()

# Carregar ambiente da raiz
load_dotenv(os.path.join(os.getcwd(), ".env"))
api_key = os.getenv("GEMINI_API_KEY")

if not api_key:
    print(Fore.RED + "ERRO FATAL: GEMINI_API_KEY não encontrada. Fechando.")
    sys.exit(1)

# Inicializações# App Setup
app = FastAPI(title="TogaMind+ Engine", version="2.0")

# -------------------------------------------------------------------
# SISTEMA DE MONITORAMENTO E TRACES DE AUDITORIA (MIDDLEWARE GLOBAL)
# -------------------------------------------------------------------
@app.middleware("http")
async def add_process_time_header(request: Request, call_next):
    start_time = time.time()
    
    # Executa a Rota
    try:
        response = await call_next(request)
        process_time = time.time() - start_time
        
        # Registra a latência e o status da rota no Arquivo Oficial
        log_info(f"ROTA: {request.method} {request.url.path} | STATUS: {response.status_code} | TEMPO: {process_time:.3f}s")
        response.headers["X-Process-Time"] = str(process_time)
        return response
    except Exception as e:
        process_time = time.time() - start_time
        registrar_erro_com_trace(e)
        raise e

# CORS Setup
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

class RAGDocument(BaseModel):
    content_type: str
    title: str
    content: str
    judge_id: str = "anonimo"

def get_active_model():
    config_path = "config.json" 
    try:
        if os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                return config.get("active_model", "gemini-3-flash")
    except Exception:
        pass
    return "gemini-3-flash"

def get_api_endpoint():
    config_path = "config.json" 
    try:
        if os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                return config.get("api_endpoint")
    except Exception:
        pass
    return None

# -------------------------------------------------------------------
# BANCO DE DADOS RELACIONAL (AUDIÊNCIAS E PAUTA DO DIA)
# -------------------------------------------------------------------
def get_db_connection():
    os.makedirs(os.path.join(os.getcwd(), "storage"), exist_ok=True)
    db_path = os.path.join(os.getcwd(), "storage", "toga_database.db")
    conn = sqlite3.connect(db_path, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db_connection()
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS audiencias (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            id_processo TEXT NOT NULL,
            data_audiencia TEXT,
            horario TEXT,
            tipo TEXT,
            partes TEXT,
            ponto_controvertido TEXT,
            status_video BOOLEAN NOT NULL DEFAULT 0,
            resumo_previa TEXT,
            lembrete TEXT,
            caminho_anexo TEXT
        )
    ''')
    conn.commit()
    conn.close()

# Inicializa o DB ao subir a API
init_db()

class AudienciaPainel(BaseModel):
    id_processo: str
    data_audiencia: Optional[str] = None
    horario: Optional[str] = None
    tipo: str
    partes: Optional[dict] = None
    ponto_controvertido: Optional[str] = None
    status_video: bool = False
    resumo_previa: Optional[str] = None
    lembrete: Optional[str] = None
    caminho_anexo: Optional[str] = None

class AudienciaResponse(AudienciaPainel):
    id: int

class AudienciaRequest(BaseModel):
    processo_id: str
    caminho_pdf: str

class AtaRequest(BaseModel):
    processo_id: str
    partes: str
    texto_ata: str
    horario_inicio: str
    horario_fim: str

import hashlib

def hash_password(password: str) -> str:
    salt = os.urandom(16).hex()
    pwd_hash = hashlib.sha256((password + salt).encode('utf-8')).hexdigest()
    return f"{salt}${pwd_hash}"

def verify_password(password: str, hashed_str: str) -> bool:
    try:
        salt, pwd_hash = hashed_str.split("$")
    except ValueError:
        return False
    return hashlib.sha256((password + salt).encode('utf-8')).hexdigest() == pwd_hash

# -------------------------------------------------------------------
# ROTAS DE AUTENTICAÇÃO E REGISTRO (MULTI-TENANCY)
# -------------------------------------------------------------------
def get_base_dir():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.getcwd()

@app.post("/register")
async def register_judge(payload: dict = Body(...)):
    judge_id = payload.get("judge_id")
    password = payload.get("password")
    if not judge_id or not password:
        raise HTTPException(status_code=400, detail="Credenciais incompletas")
        
    base_dir = get_base_dir()
    vault_path = os.path.join(base_dir, "storage", "rag_vault", judge_id)
    
    try:
        os.makedirs(vault_path, exist_ok=True)
    except Exception as e:
        logger.error(f"Erro ao criar diretorio {vault_path}: {e}")
        raise HTTPException(status_code=500, detail=f"Erro de I/O de disco: {e}")
    
    config_file = os.path.join(vault_path, "user_config.json")
    if os.path.exists(config_file):
        raise HTTPException(status_code=400, detail="Magistrado já cadastrado.")
        
    pwd_hashed = hash_password(password)
    try:
        with open(config_file, "w", encoding="utf-8") as f:
            json.dump({"password_hash": pwd_hashed}, f)
    except Exception as e:
        logger.error(f"Erro ao salvar config: {e}")
        raise HTTPException(status_code=500, detail=f"Erro de escrita: {e}")
        
    return {"status": "success", "message": "Gabinete criado com sucesso!"}

@app.post("/login")
async def login_judge(payload: dict = Body(...)):
    judge_id = payload.get("judge_id")
    password = payload.get("password")
    if not judge_id or not password:
        raise HTTPException(status_code=400, detail="Credenciais incompletas")
        
    base_dir = get_base_dir()
    config_file = os.path.join(base_dir, "storage", "rag_vault", judge_id, "user_config.json")
    
    if not os.path.exists(config_file):
        raise HTTPException(status_code=401, detail="Magistrado não encontrado.")
        
    try:
        with open(config_file, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        logger.error(f"Erro lendo o arquivo config {config_file}: {e}")
        raise HTTPException(status_code=500, detail="Erro interno lendo os dados")
        
    hashed_str = data.get("password_hash")
    if not hashed_str or not verify_password(password, hashed_str):
        raise HTTPException(status_code=401, detail="Senha incorreta.")
        
    return {"status": "success", "message": "Autenticado"}

# -------------------------------------------------------------------
# ROTAS DE IA E RAG
# -------------------------------------------------------------------

@app.post("/analyze")
async def analyze_process(
    request: Request,
    file: UploadFile = File(...)
):
    headers = dict(request.headers)
    logger.info(f"[/analyze] START: Requisição recebida.")
    logger.info(f"[/analyze] HEADERS CRUS: {headers}")
    
    judge_id = headers.get("judge-id", headers.get("judge_id", "anonimo"))

    if not judge_id or judge_id == "null":
        judge_id = "anonimo"
        
    logger.info(f"[/analyze] judge_id final: {judge_id}")

    try:
        pdf_content = await file.read()
        logger.info(f"[/analyze] PDF lido com sucesso. ({len(pdf_content)} bytes).")
        
        active_model = get_active_model()
        logger.info(f"[/analyze] Modelo ativo selecionado: {active_model}")
        
        pdf_b64 = base64.b64encode(pdf_content).decode('utf-8')
        
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{active_model}:generateContent?key={api_key}"
        
        headers = {'Content-Type': 'application/json'}
        
        # Preparação do Histórico Vivo do Gabinete (Zero Cloud)
        from toga_rag_memory import carregar_contexto_historico
        memoria_resultado = carregar_contexto_historico(judge_id) # Pass judge_id to filter memory
        historico_contexto = ""
        usando_historico = False
        
        if isinstance(memoria_resultado, dict) and "historico_carregado" in memoria_resultado and memoria_resultado["historico_carregado"] > 0:
            usando_historico = True
            historico_contexto = "\n--- HISTÓRICO DE DECISÕES DO GABINETE ---\n" + "\n".join(memoria_resultado["dados"][:3])
        
        system_instruction_base = PROMPT_MESTRE_BASE + (
            "\nTAREFA DE ANÁLISE:\n"
            "Analise os documentos fornecidos via RAG e elabore um parecer técnico.\n"
            "OBRIGATÓRIO: Formate toda a sua resposta utilizando MARKDOWN ESTRITO. "
            "Sempre que houver comparação de valores, datas ou versões de partes diferentes, utilize Tabelas Markdown com as colunas: | Item | Versão Autor | Versão Réu | Fonte |.\n\n"
            "Sua resposta deve seguir obrigatoriamente este formato Markdown:\n"
            "📋 Resumo da Análise\n"
            "[Breve parágrafo]\n\n"
            "🔍 Pontos Chave\n"
            "Fato: [Descrição] | Fonte: [Documento/Página]\n\n"
            "⚖️ Fundamentação Legal\n"
            "[Citação da Lei ou Jurisprudência aplicável baseada APENAS no contexto]\n\n"
            "Conclusão: [Proposta de encaminhamento ou decisão técnica]"
        )

        # Augment system instruction with historical context if available
        if usando_historico:
            system_instruction = (
                f"Você é o TogaMind+, Assistente Jurídico de Gabinete (Modelo Offline Strict).\n"
                f"Aja de forma neutra, culta e protocolar.\n\n"
                f"Abaixo está o texto cru de uma peça jurídica e o histórico de julgamentos desse Juiz:\n"
                f"{historico_contexto}\n\n"
                f"{system_instruction_base}"
            )
        else:
            system_instruction = system_instruction_base
        
        payload = {
            "system_instruction": {
                "parts": [{"text": system_instruction}]
            },
            "contents": [
                {
                    "parts": [
                        {"text": "Analise este processo judicial e destaque os pontos críticos para o magistrado:"},
                        {
                            "inline_data": {
                                "mime_type": "application/pdf",
                                "data": pdf_b64
                            }
                        }
                    ]
                }
            ],
            "generationConfig": {
                "candidateCount": 1,
                "temperature": 0.0
            }
        }
        
        logger.info(f"[/analyze] Preparando Payload POST para Rest API do Gemini.")
        
        response = requests.post(url, headers=headers, json=payload)
        
        logger.info(f"[/analyze] Resposta do Gemini recebida (Status: {response.status_code}).")
        
        if response.status_code != 200:
            logger.error(f"[/analyze] Erro na Resposta da API: {response.text}")
            raise Exception(f"Erro na API Gemini: {response.text}")
            
        data = response.json()
        analysis_text = data['candidates'][0]['content']['parts'][0]['text']
        
        get_rag_manager().save_for_rag(
            content_type="peticao_analisada",
            title=file.filename or "processo",
            content=analysis_text,
            judge_id=judge_id
        )

        return {"analysis": analysis_text, "model_used": active_model}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/ask-toga")
async def ask_toga(
    request: Request,
    payload: dict = Body(...)
):
    judge_id = request.headers.get("judge-id", request.headers.get("judge_id"))
    if not judge_id:
        raise HTTPException(status_code=401, detail="Acesso Negado. Credenciais ausentes.")

    query = payload.get("query")
    if not query:
        raise HTTPException(status_code=400, detail="Pergunta inválida.")

    try:
        # Recupera os contextos arquivados isolados por judge_id no drive local
        context_docs = get_vector_engine().search_similar(query, top_k=3, judge_id=judge_id)
        context_text = "\n\n".join(context_docs)

        active_model = get_active_model()
        api_endpoint = get_api_endpoint()
        
        import google.generativeai as genai
        if api_endpoint:
            genai.configure(api_key=api_key, client_options={'api_endpoint': api_endpoint})
        else:
            genai.configure(api_key=api_key)
        
        model = genai.GenerativeModel(
            model_name=active_model,
            system_instruction=(
                "[INSTRUÇÃO DE CONFINAMENTO CRÍTICA]\n"
                "Você é um assistente de análise jurídica que opera sob o protocolo de Alucinação Zero. "
                "Abaixo, você receberá trechos recuperados de documentos processuais. Sua resposta deve ser construída exclusivamente com base nesses dados.\n\n"
                "REGRAS DE OURO:\n"
                "FIDELIDADE AOS AUTOS: Se a informação solicitada pelo juiz não estiver presente nos trechos abaixo, você deve responder exatamente: \"Informação não localizada nos documentos do processo\". É terminantemente proibido usar seu conhecimento geral para preencher lacunas.\n"
                "CITAÇÃO OBRIGATÓRIA: Toda e qualquer afirmação deve ser seguida do nome do documento e página/trecho de onde foi extraída. Ex: \"O réu alega prescrição (Contestação, fls. 45)\".\n"
                "PROIBIÇÃO DE INVENÇÃO: Não invente números de artigos, leis ou precedentes. Cite apenas o que está no texto ou leis federais brasileiras se forem explicitamente mencionadas.\n"
                "DIVERGÊNCIAS: Se o Documento A diz algo diferente do Documento B, reporte o conflito: \"Há divergência entre a Inicial [Doc 1] e a Perícia [Doc 3] quanto ao valor do dano\"."
            ),
            generation_config={"candidate_count": 1, "temperature": 0.0}
        )

        full_prompt = (
            f"CONTEXTO RECUPERADO (FONTE DA VERDADE):\n{context_text}\n\n"
            f"PERGUNTA DO MAGISTRADO:\n{query}\n\n"
            "RESPOSTA TÉCNICA (FORMATO DE NOTA DE GABINETE):"
        )

        response = model.generate_content(full_prompt)

        return {
            "answer": response.text,
            "used_context": len(context_docs) > 0,
            "model": active_model
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/process_pdf")
async def get_process_pdf(path: str):
    if not os.path.exists(path) or not path.lower().endswith(".pdf"):
        raise HTTPException(status_code=404, detail="Arquivo PDF não encontrado.")
    return FileResponse(path, media_type="application/pdf")

@app.post("/chat-contextual")
async def chat_contextual(
    request: Request,
    payload: dict = Body(...)
):
    judge_id = request.headers.get("judge-id", request.headers.get("judge_id"))
    if not judge_id:
        raise HTTPException(status_code=401, detail="Acesso Negado. Credenciais ausentes.")

    query = payload.get("query")
    processo_numero = payload.get("processo_numero")
    
    if not query or not processo_numero:
        raise HTTPException(status_code=400, detail="Pergunta ou número do processo inválido.")

    # 1. Definir o escopo da busca no Drive E:
    caminho_processo = os.path.join(os.getcwd(), "storage", "rag_vault", judge_id, "processos", f"{processo_numero}.pdf")
    
    if not os.path.exists(caminho_processo):
         raise HTTPException(status_code=404, detail="Autos do processo não encontrados na base local.")

    try:
        active_model = get_active_model()
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        
        model = genai.GenerativeModel(
            model_name=active_model,
            system_instruction=(
                PROMPT_MESTRE_BASE +
                "\nATUAÇÃO ESPECÍFICA (CHAT DE AUTOS):\n"
                "Responda à pergunta do usuário baseando-se EXCLUSIVAMENTE nos autos do processo fornecido no anexo PDF. "
                "CITE A PÁGINA do documento em sua resposta obrigatoriamente, se aplicável. "
                "Retorne os dados estritamente no esquema: {'texto': 'Sua resposta aqui.', 'pagina': 12}. "
                "Se não encontrar o número da página exato, retorne -1 no campo pagina."
            ),
            generation_config={
                "temperature": 0.0,
                "response_mime_type": "application/json",
            }
        )

        with open(caminho_processo, "rb") as f:
            pdf_data = f.read()

        full_prompt = (
            f"PERGUNTA ATUAL referente a este processo: {query}\n\n"
            "Responda de forma técnica, inserindo o número da página no campo 'pagina'."
        )

        response = model.generate_content([
            full_prompt,
            {'mime_type': 'application/pdf', 'data': pdf_data}
        ])
        
        # Deserializar a resposta assegurada pelo JSON MimeType
        response_data = json.loads(response.text)

        return {
            "resposta": response_data.get("texto", "Não foi possível estruturar a resposta."),
            "pagina": response_data.get("pagina", -1),
            "arquivo": caminho_processo,
            "model": active_model
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/gerar-minuta")
async def gerar_minuta(
    request: Request,
    payload: dict = Body(...)
):
    judge_id = request.headers.get("judge-id", request.headers.get("judge_id"))
    if not judge_id:
        raise HTTPException(status_code=401, detail="Acesso Negado. Credenciais ausentes.")

    ponto_decisao = payload.get("ponto_decisao")
    processo_numero = payload.get("processo_numero")

    if not ponto_decisao or not processo_numero:
        raise HTTPException(status_code=400, detail="Ponto ou processo inválido.")

    # 1. Recupera as citações e páginas já indexadas no Drive E:
    caminho_processo = os.path.join(os.getcwd(), "storage", "rag_vault", judge_id, "processos", f"{processo_numero}.pdf")
    if not os.path.exists(caminho_processo):
         raise HTTPException(status_code=404, detail="Autos do processo não encontrados na base local.")

    try:
        active_model = get_active_model()
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        
        model = genai.GenerativeModel(
            model_name=active_model,
            system_instruction=(
                PROMPT_MESTRE_BASE +
                "\nATUAÇÃO ESPECÍFICA (REDAÇÃO DE MINUTA):\n"
                "Redija parágrafos de fundamentação (minutas de decisão) baseadas unicamente na evidência do anexo PDF. "
                "Regra: Use citação direta indicando a página do PDF entre parênteses (ex: fls. X). "
                "Estilo: Sóbrio, impessoal e estritamente técnico e jurídico.\n\n"
                "IMPORTANTE PARA EXPORTAÇÃO PDF: Gere o relatório em texto limpo, use apenas negrito para títulos e evite o "
                "uso de tabelas complexas ou emojis dentro do corpo do relatório, pois este texto será convertido diretamente em um documento oficial PDF."
            ),
            generation_config={"temperature": 0.0}
        )

        with open(caminho_processo, "rb") as f:
            pdf_data = f.read()

        full_prompt = (
            f"Com base nos autos {processo_numero}, redija um parágrafo de fundamentação "
            f"sobre '{ponto_decisao}'."
        )

        response = model.generate_content([
            full_prompt,
            {'mime_type': 'application/pdf', 'data': pdf_data}
        ])

        return {"minuta": response.text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/exportar-pdf-decisao")
async def exportar_pdf(
    request: Request,
    payload: dict = Body(...)
):
    judge_id = request.headers.get("judge-id", request.headers.get("judge_id"))
    if not judge_id:
        raise HTTPException(status_code=401, detail="Acesso Negado. Credenciais ausentes.")

    conteudo = payload.get("conteudo")
    processo_numero = payload.get("processo_numero")

    if not conteudo or not processo_numero:
        raise HTTPException(status_code=400, detail="Conteúdo ou número do processo inválido.")

    # 1. Definir caminho de saída no Drive E:
    output_path = os.path.join(os.getcwd(), "storage", "rag_vault", judge_id, "decisoes", f"Decisao_{processo_numero}.pdf")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # 2. Criar o PDF com Padrão 2026
    c = canvas.Canvas(output_path, pagesize=A4)
    c.setFont("Helvetica-Bold", 12)
    c.drawString(100, 800, "PODER JUDICIÁRIO - ESTADO DE SÃO PAULO")
    
    c.setFont("Helvetica", 11)
    # Quebra de linha automática (Primitiva) para a minuta gerada
    textobject = c.beginText(100, 750)
    for line in conteudo.split('\n'):
        textobject.textLine(line)
        # Handle long lines primitively by wrapping if extremely long (Simpler implementation given ReportLab canvas.textLine restrictions)
        
    c.drawText(textobject)
    
    # 3. Rodapé Obrigatório (Protocolo 2026)
    c.setFont("Helvetica-Oblique", 8)
    footer_text = "Página 1 | © 2026 ScanNut Multiverso Digital"
    c.drawString(200, 50, footer_text)
    
    c.save()
    return {"status": "success", "file_url": output_path}

@app.post("/save-rag")
async def save_to_rag(
    request: Request,
    doc: RAGDocument
):
    judge_id = request.headers.get("judge-id", request.headers.get("judge_id"))
    if not judge_id:
        raise HTTPException(status_code=401, detail="Acesso Negado.")
        
    try:
        saved_path = get_rag_manager().save_for_rag(
            content_type=doc.content_type,
            title=doc.title,
            content=doc.content,
            judge_id=judge_id
        )
        return {"status": "success", "saved_path": saved_path}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/sync-rag")
async def sync_rag(
    request: Request,
    payload: dict = Body(...)
):
    judge_id = request.headers.get("judge-id", request.headers.get("judge_id"))
    if not judge_id:
        raise HTTPException(status_code=401, detail="Acesso Negado.")

    text = payload.get("content")
    if text:
        get_vector_engine().add_document(text, judge_id=judge_id)
        return {"status": "Sincronizado com a memória RAG"}
    return {"status": "Conteúdo vazio"}

@app.get("/config.json")
async def get_config():
    config_path = "config.json"
    if os.path.exists(config_path):
        return FileResponse(config_path)
    raise HTTPException(status_code=404, detail="config.json não encontrado")

import shutil
from services.court_integration import download_processo_com_token

active_sessions = {}

@app.get("/token-status")
async def token_status(request: Request):
    judge_id = request.headers.get("judge-id", request.headers.get("judge_id"))
    if not judge_id:
        raise HTTPException(status_code=401, detail="Usuário não identificado.")
    pfx_path = os.path.join(os.getcwd(), "storage", "rag_vault", judge_id, "credentials", "token_magistrado.pfx")
    registered = os.path.exists(pfx_path)
    unlocked = judge_id in active_sessions
    return {"registered": registered, "unlocked": unlocked}

@app.post("/register-token")
async def register_token(
    request: Request,
    certificate: UploadFile = File(...)
):
    judge_id = request.headers.get("judge-id", request.headers.get("judge_id"))
    if not judge_id:
        raise HTTPException(status_code=401, detail="Usuário não identificado.")
    
    target_dir = os.path.join(os.getcwd(), "storage", "rag_vault", judge_id, "credentials")
    os.makedirs(target_dir, exist_ok=True)
    file_path = os.path.join(target_dir, "token_magistrado.pfx")
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(certificate.file, buffer)
        
    return {"status": "success", "path": file_path}

@app.post("/unlock-token")
async def unlock_token(
    request: Request,
    password: str = Body(embed=True)
):
    judge_id = request.headers.get("judge-id", request.headers.get("judge_id"))
    if not judge_id:
        raise HTTPException(status_code=401, detail="Usuário não identificado.")

    pfx_path = os.path.join(os.getcwd(), "storage", "rag_vault", judge_id, "credentials", "token_magistrado.pfx")
    
    if not os.path.exists(pfx_path):
        return {"status": "error", "message": "Token não encontrado no gabinete."}

    try:
        from cryptography.hazmat.primitives.serialization import pkcs12
        from cryptography.hazmat.backends import default_backend
        with open(pfx_path, "rb") as f:
            pkcs12.load_key_and_certificates(f.read(), password.encode(), default_backend())
        
        active_sessions[judge_id] = password
        return {"status": "success", "message": "Token desbloqueado com sucesso."}
    except Exception:
        return {"status": "error", "message": "Senha do certificado incorreta."}

@app.post("/import-process")
async def import_process(
    request: Request,
    process_number: str = Form(...)
):
    judge_id = request.headers.get("judge-id", request.headers.get("judge_id"))
    if not judge_id:
        raise HTTPException(status_code=401, detail="Acesso Negado.")

    if judge_id not in active_sessions:
        raise HTTPException(status_code=401, detail="Sessão de token não validada.")

    pfx_path = os.path.join(os.getcwd(), "storage", "rag_vault", judge_id, "credentials", "token_magistrado.pfx")
    password = active_sessions[judge_id]

    try:
        success = download_processo_com_token(judge_id, pfx_path, password, process_number)
        
        if success:
            pdf_dest = os.path.join(os.getcwd(), "storage", "rag_vault", judge_id, "processos", f"{process_number}.pdf")
            resumo = await gerar_resumo_imediato(pdf_dest)
            
            return {
                "status": "success", 
                "message": "Autos importados com sucesso na sua base pessoal.",
                "summary": resumo
            }
        else:
            raise HTTPException(status_code=500, detail="Falha ao baixar autos com o token fornecido.")
                
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro de processamento do token: {str(e)}")

# -------------------------------------------------------------------
# ROTAS CRUD - PAUTA DE AUDIÊNCIAS (PROTOCOLO SCANNUT)
# -------------------------------------------------------------------

@app.post("/sync-audiencias")
async def sync_pauta_audiencias(request: Request):
    judge_id = request.headers.get("judge-id", request.headers.get("judge_id", "anonimo"))
    
    # 1. Definir pasta do Juiz no Vault
    pasta_processos = os.path.join(os.getcwd(), "storage", "rag_vault", judge_id, "processos")
    
    if not os.path.exists(pasta_processos):
        return {"status": "error", "message": "Nenhum processo salvo pelo Gabinete."}
        
    try:
        from toga_audiencia_scanner import TogaAudienciaScanner
        scanner = TogaAudienciaScanner()
        active_model = get_active_model()
        
        resultado = scanner.varrer_processos(
            judge_id=judge_id,
            pasta_processos=pasta_processos,
            api_key=api_key,
            active_model=active_model
        )
        return resultado
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/audiencia/preparar")
async def preparar_audiencia(req: AudienciaRequest):
    try:
        if not os.path.exists(req.caminho_pdf):
            raise HTTPException(status_code=404, detail="Arquivo PDF não encontrado no caminho especificado.")

        import google.generativeai as genai
        genai.configure(api_key=api_key)
        active_model = get_active_model()
        model = genai.GenerativeModel(active_model)

        from toga_pdf_reader import TogaPDFReader
        extracao_bruta = TogaPDFReader.extrair_resumo_peticao(req.caminho_pdf, max_pages=10)
        
        texto_base = extracao_bruta.get("resumo_fato", "") + "\n\n" + extracao_bruta.get("objetivo_prova", "")
        if "erro" in extracao_bruta:
             texto_base = "Não foi possível realizar a pré-extração Nativa."

        prompt = f"""
        CONTEXTO DE GABINETE:
        Analise o processo {req.processo_id} focado no documento inserido.
        Extraia do texto abaixo os pontos cruciais para a audiência de hoje:
        1. RESUMO DOS FATOS.
        2. PONTOS CONTROVERTIDOS (O que precisa de prova).
        3. SUGESTÕES DE PERGUNTAS (Para o magistrado).
        
        TEXTO: {texto_base[:2500]}
        """
        
        response = model.generate_content(prompt)
        
        # O RAG preenche automaticamente o DB para otimizar futuras exibições
        conn = get_db_connection()
        c = conn.cursor()
        c.execute('''
            UPDATE audiencias SET ponto_controvertido = ? WHERE id_processo = ?
        ''', (response.text, req.processo_id))
        conn.commit()
        conn.close()
        
        return {
            "status": "success",
            "processo": req.processo_id,
            "analise_ia": response.text,
            "historico_utilizado": usando_historico
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/gerar-ata-docx")
async def gerar_ata_docx(req: AtaRequest):
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import datetime

        # 1. Configurar Pasta de Destino (Relatórios do Gabinete na Área de Trabalho ou Pasta Raiz)
        pasta_atas = os.path.expanduser("~/Desktop/Relatorios_Assistente")
        os.makedirs(pasta_atas, exist_ok=True)
        
        # 2. Construir Documento Vazio (Padrão Judiciário Básico)
        doc = Document()
        
        # Cabeçalho
        h1 = doc.add_heading(f'TERMO DE AUDIÊNCIA - PROCESSO {req.processo_id}', 1)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadados
        doc.add_paragraph(f"Partes: {req.partes}", style='Normal')
        doc.add_paragraph(f"Data Base: {datetime.datetime.now().strftime('%d/%m/%Y')} | Início: {req.horario_inicio} - Término: {req.horario_fim}", style='Normal')
        
        doc.add_paragraph("-" * 80)
        
        # Corpo da Ata
        doc.add_heading('OCORRÊNCIAS / CELEBRAÇÃO', level=2)
        p = doc.add_paragraph(req.texto_ata)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Inserção do Rodapé Padrão
        try:
            section = doc.sections[0]
            footer = section.footer
            footer.paragraphs[0].text = "© 2026 TogaMind+ | Gabinete Digital Local"
        except Exception:
            pass # Prevenção de quebra de docx layout
        
        # 3. Salvar Fisicamente (Módulo Pós-Audiência)
        # O arquivo nasce direto na pasta de assinatura.
        pasta_assinatura = os.path.join(pasta_atas, "Para_Assinar")
        os.makedirs(pasta_assinatura, exist_ok=True)
        
        nome_arquivo = f"Ata_Finalizada_{req.processo_id.replace('.', '_').replace('-', '_')}.docx"
        caminho_final = os.path.join(pasta_assinatura, nome_arquivo)
        
        doc.save(caminho_final)
        
        # 4. Feedback Local: Remove Audiência da Pauta ativa
        conn = get_db_connection()
        c = conn.cursor()
        c.execute('DELETE FROM audiencias WHERE id_processo = ?', (req.processo_id,))
        conn.commit()
        conn.close()

        return {
            "status": "success",
            "mensagem": "Ata gerada e salva com sucesso.",
            "caminho_salvo": caminho_final
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/backup")
async def realizar_backup_seguranca():
    try:
        from toga_backup_service import executar_backup_gabinete
        resultado = executar_backup_gabinete()
        if "trace" in resultado:
             raise Exception(resultado["message"])
        return resultado
    except Exception as e:
         raise HTTPException(status_code=500, detail=str(e))

@app.get("/audiencias", response_model=List[AudienciaResponse])
async def listar_audiencias():
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("SELECT * FROM audiencias ORDER BY horario ASC")
    rows = c.fetchall()
    conn.close()
    
    resultados = []
    for r in rows:
        d = dict(r)
        if d.get("partes"):
            try:
                d["partes"] = json.loads(d["partes"])
            except Exception:
                d["partes"] = {}
        resultados.append(d)
        
    return resultados

@app.post("/audiencias", response_model=AudienciaResponse)
async def criar_audiencia(audiencia: AudienciaPainel):
    conn = get_db_connection()
    c = conn.cursor()
    
    partes_str = json.dumps(audiencia.partes) if audiencia.partes else "{}"
    
    c.execute('''
        INSERT INTO audiencias (id_processo, data_audiencia, horario, tipo, partes, ponto_controvertido, status_video, resumo_previa, lembrete, caminho_anexo)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        audiencia.id_processo, audiencia.data_audiencia, audiencia.horario, audiencia.tipo,
        partes_str, audiencia.ponto_controvertido,
        int(audiencia.status_video), audiencia.resumo_previa, audiencia.lembrete, audiencia.caminho_anexo
    ))
    conn.commit()
    novo_id = c.lastrowid
    conn.close()
    
    return {**audiencia.dict(), "id": novo_id}

@app.delete("/audiencias/{audiencia_id}")
async def deletar_audiencia(audiencia_id: int):
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("DELETE FROM audiencias WHERE id = ?", (audiencia_id,))
    conn.commit()
    if c.rowcount == 0:
        conn.close()
        raise HTTPException(status_code=404, detail="Audiência não encontrada.")
    conn.close()
    return {"status": "success", "message": "Audiência deletada do sistema."}

async def gerar_resumo_imediato(pdf_path: str) -> str:
    prompt = """
    Analise os autos recém-capturados e forneça um resumo em 3 pontos:
    1. Objeto principal da ação.
    2. Última movimentação relevante (decisão ou petição).
    3. Pendência imediata (prazo ou conclusão).
    Seja conciso e use linguagem jurídica sóbria.
    """
    try:
        active_model = get_active_model()
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(
            model_name=active_model,
            generation_config={"candidate_count": 1, "temperature": 0.2}
        )
        with open(pdf_path, "rb") as f:
            pdf_data = f.read()
            
        response = model.generate_content([
            prompt,
            {'mime_type': 'application/pdf', 'data': pdf_data}
        ])
        return response.text
    except Exception as e:
        return f"Não foi possível gerar o resumo automático: {str(e)}"

# -------------------------------------------------------------------
# MONTAGEM DA INTERFACE WEB ESTÁTICA (DEVE SER O ÚLTIMO)
# -------------------------------------------------------------------
WEB_PATH = os.path.join(getattr(sys, '_MEIPASS', os.getcwd()), "build", "web")
if os.path.exists(WEB_PATH):
    app.mount("/", StaticFiles(directory=WEB_PATH, html=True), name="ui")
else:
    fallback_path = r"E:\antigravity_projetos\toga_mind_plus\build\web"
    if os.path.exists(fallback_path):
        app.mount("/", StaticFiles(directory=fallback_path, html=True), name="ui")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
